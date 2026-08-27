import dataclasses
import shutil
from pathlib import Path

import docx
import pytest

import excel_reader
import paths
import section_dao_duc
import tokens
import word_writer

CHECKLIST_PATH = paths.project_root() / excel_reader.CHECKLIST_FILENAME
SHEET_VIAM = "Đề tài - Bánh ăn dặm VIAM 2027"

SOURCE_DIR = paths.project_root() / "01. Hồ sơ đạo đức đề cương - MẪU"
FILES = [
    "00. QĐ Giao đề tài.docx",
    "01. QĐTLHĐ đạo đức đề cương.docx",
    "02. BB họp HĐ đạo đức.docx",
    "03. BB kiểm phiếu HĐ đạo đức.docx",
    "04. QĐ chấp nhận đạo đức.docx",
    "Bảng kiểm đánh giá đạo đức.docx",
]


@pytest.fixture()
def dest_dir(tmp_path):
    for filename in FILES:
        shutil.copy2(SOURCE_DIR / filename, tmp_path / filename)
    return tmp_path


@pytest.fixture()
def info():
    return excel_reader.load_project_data(CHECKLIST_PATH, SHEET_VIAM)


def _available_backends():
    backends = ["docx"]
    if word_writer.com_available():
        backends.append("com")
    return backends


BACKENDS = _available_backends()


def test_generate_fixes_ethics_secretary_org(dest_dir, info):
    session = word_writer.Session(force_backend="docx")
    try:
        section_dao_duc.generate(session, dest_dir, info, tokens.build_common_tokens(info))
    finally:
        session.quit()

    doc = docx.Document(str(dest_dir / "01. QĐTLHĐ đạo đức đề cương.docx"))
    secretary_table = doc.tables[2]
    assert secretary_table.cell(0, 0).text.strip() == "Hoàng Hà Linh"
    assert secretary_table.cell(0, 1).text.strip() == "Trung tâm NCKH - Viện VIAM"


@pytest.mark.parametrize("backend", BACKENDS)
def test_generate_replaces_title_everywhere(dest_dir, info, backend):
    session = word_writer.Session(force_backend=backend)
    try:
        section_dao_duc.generate(session, dest_dir, info, tokens.build_common_tokens(info))
    finally:
        session.quit()

    for filename in ["00. QĐ Giao đề tài.docx", "03. BB kiểm phiếu HĐ đạo đức.docx"]:
        doc = docx.Document(str(dest_dir / filename))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert info.title in full_text
        assert "{{TEN_DE_TAI}}" not in full_text


def test_generate_writes_head_and_researchers_into_giao_de_tai(dest_dir, info):
    session = word_writer.Session(force_backend="docx")
    try:
        section_dao_duc.generate(session, dest_dir, info, tokens.build_common_tokens(info))
    finally:
        session.quit()

    doc = docx.Document(str(dest_dir / "00. QĐ Giao đề tài.docx"))
    unit_table = doc.tables[2]
    head_cell_text = unit_table.cell(1, 2).text
    members_cell_text = unit_table.cell(2, 2).text

    assert f"{info.head.degree} {info.head.name}".strip() in head_cell_text
    assert "Cử nhân HOÀNG HÀ LINH" not in members_cell_text
    for researcher in info.researchers:
        assert f"{researcher.degree} {researcher.name}".strip() in members_cell_text


def test_generate_with_word_com_writes_head_without_raising(dest_dir, info):
    if not word_writer.com_available():
        pytest.skip("Can Word COM de chay test nay")

    session = word_writer.Session(force_backend="com")
    try:
        section_dao_duc.generate(session, dest_dir, info, tokens.build_common_tokens(info))
    finally:
        session.quit()

    doc = docx.Document(str(dest_dir / "00. QĐ Giao đề tài.docx"))
    unit_table = doc.tables[2]
    head_cell_text = unit_table.cell(1, 2).text
    assert f"{info.head.degree} {info.head.name}".strip() in head_cell_text


def test_generate_uses_parsed_timeline_not_just_year(dest_dir, info):
    custom_info = dataclasses.replace(info, timeline="Tháng 03/2027 đến tháng 09/2028")
    session = word_writer.Session(force_backend="docx")
    try:
        section_dao_duc.generate(session, dest_dir, custom_info, tokens.build_common_tokens(custom_info))
    finally:
        session.quit()

    doc = docx.Document(str(dest_dir / "04. QĐ chấp nhận đạo đức.docx"))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Từ 03/2027 đến 09/2028" in full_text

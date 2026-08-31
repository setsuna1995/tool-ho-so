import docx
import pytest
from pathlib import Path

import docx_template_engine as engine


def _create_simple_template(path: Path, text: str):
    doc = docx.Document()
    doc.add_paragraph(text)
    doc.save(path)


def test_is_jinja_template(tmp_path):
    p_jinja = tmp_path / "jinja.docx"
    _create_simple_template(p_jinja, "Đề tài: {{ project.title }}")

    p_plain = tmp_path / "plain.docx"
    _create_simple_template(p_plain, "Văn bản thông thường không có tag")

    assert engine.is_jinja_template(p_jinja) is True
    assert engine.is_jinja_template(p_plain) is False


def test_render_template_simple_and_filter(tmp_path):
    tmpl_path = tmp_path / "template.docx"
    _create_simple_template(tmpl_path, "Tên: {{ project.title | upper }} - Năm: {{ project.year }}")

    out_path = tmp_path / "output.docx"
    context = {
        "project": {
            "title": "Nghiên cứu bánh ăn dặm",
            "year": 2027,
        }
    }
    engine.render_template(tmpl_path, out_path, context)

    res_doc = docx.Document(out_path)
    full_text = "\n".join(p.text for p in res_doc.paragraphs)
    assert "NGHIÊN CỨU BÁNH ĂN DẶM" in full_text
    assert "2027" in full_text


def test_render_template_condition_and_loop(tmp_path):
    tmpl_path = tmp_path / "template_loop.docx"
    doc = docx.Document()
    doc.add_paragraph("{% if project.has_partner %}Đối tác: {{ project.partner }}{% endif %}")
    doc.add_paragraph("{% for r in researchers %}{{ loop.index }}. {{ r.name }}\n{% endfor %}")
    doc.save(tmpl_path)

    out_path = tmp_path / "output_loop.docx"
    context = {
        "project": {
            "has_partner": False,
            "partner": "Không hiện",
        },
        "researchers": [
            {"name": "Nguyễn Văn A"},
            {"name": "Trần Thị B"},
        ],
    }
    engine.render_template(tmpl_path, out_path, context)

    res_doc = docx.Document(out_path)
    full_text = "\n".join(p.text for p in res_doc.paragraphs)
    assert "Đối tác:" not in full_text
    assert "1. Nguyễn Văn A" in full_text
    assert "2. Trần Thị B" in full_text

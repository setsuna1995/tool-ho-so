# test_expert_invitation.py
import docx
from docx.oxml.ns import qn

import expert_invitation
from excel_reader import Person


def _make_template(tmp_path):
    path = tmp_path / "letter.docx"
    doc = docx.Document()
    doc.add_paragraph("Đề tài: {{TEN_DE_TAI}}")
    doc.add_paragraph("Kính gửi: {{CHUYEN_GIA_HO_TEN}}, {{CHUYEN_GIA_DON_VI}}.")
    doc.save(str(path))
    return path


def test_returns_false_and_leaves_file_untouched_when_no_recipients(tmp_path):
    path = _make_template(tmp_path)
    original_text = [p.text for p in docx.Document(str(path)).paragraphs]

    result = expert_invitation.generate_multi_page_letter(path, [], {"{{TEN_DE_TAI}}": "Đề tài X"})

    assert result is False
    assert [p.text for p in docx.Document(str(path)).paragraphs] == original_text


def test_generates_one_page_per_recipient_with_correct_names(tmp_path):
    path = _make_template(tmp_path)
    recipients = [
        Person("Hoàng Thị Thanh", "PGs. Ts.", "Hội đồng Đạo đức trong nghiên cứu y sinh học Quốc gia"),
        Person("Nguyễn Công Khẩn", "Gs. Ts.", "Hiệp hội Sữa Việt Nam"),
    ]

    result = expert_invitation.generate_multi_page_letter(path, recipients, {"{{TEN_DE_TAI}}": "Đề tài X"})

    assert result is True
    doc = docx.Document(str(path))
    full_text = [p.text for p in doc.paragraphs]
    assert full_text.count("Đề tài: Đề tài X") == 2
    assert "Kính gửi: PGs. Ts. Hoàng Thị Thanh, Hội đồng Đạo đức trong nghiên cứu y sinh học Quốc gia." in full_text
    assert "Kính gửi: Gs. Ts. Nguyễn Công Khẩn, Hiệp hội Sữa Việt Nam." in full_text
    assert "{{CHUYEN_GIA_HO_TEN}}" not in "\n".join(full_text)


def test_inserts_exactly_one_page_break_between_two_recipients(tmp_path):
    path = _make_template(tmp_path)  # template has 2 paragraphs
    recipients = [Person("Người Một", "", "Đơn vị 1"), Person("Người Hai", "", "Đơn vị 2")]

    expert_invitation.generate_multi_page_letter(path, recipients, {"{{TEN_DE_TAI}}": "X"})

    doc = docx.Document(str(path))
    assert len(doc.paragraphs) == 2 * 2 + 1  # 2 template paragraphs x 2 pages + 1 page-break paragraph


def test_sect_pr_remains_last_body_element(tmp_path):
    path = _make_template(tmp_path)
    recipients = [Person("Người Một", "", "Đơn vị 1")]

    expert_invitation.generate_multi_page_letter(path, recipients, {})

    doc = docx.Document(str(path))
    assert doc.element.body[-1].tag == qn("w:sectPr")

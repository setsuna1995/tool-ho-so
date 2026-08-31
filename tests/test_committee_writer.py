import docx
import pytest

import committee_writer
import word_writer
from excel_reader import CommitteeData, Person


def _sample_committee():
    return CommitteeData(
        chair=Person("Nguyễn Công Khẩn", "GS.TS.", "Hội đồng Đạo đức Y sinh Quốc gia"),
        reviewers=[Person("Nguyễn Xuân Ninh", "PGS.TS.BS.", "Viện Y học Ứng dụng Việt Nam")],
        members=[Person("Nguyễn Thị Lâm", "PGS.TS.BS.", "Viện Dinh dưỡng QG")],
        secretaries=[
            Person("Hoàng Hà Linh", "Thạc sĩ", "Trung tâm NCKH - Viện VIAM"),
            Person("Trương Phan Hồng Hà", "Thạc sĩ", "Viện Y học Ứng dụng Việt Nam"),
        ],
    )


def _table_fixture(tmp_path, rows, cols):
    path = tmp_path / "table.docx"
    d = docx.Document()
    d.add_table(rows=rows, cols=cols)
    d.save(str(path))
    return path


def test_write_committee_roster_fills_name_org_role(tmp_path):
    src = _table_fixture(tmp_path, rows=3, cols=3)
    session = word_writer.Session(force_backend="docx")
    try:
        doc = session.open(src)
        committee = _sample_committee()
        committee_writer.write_committee_roster(
            session, doc, 1, committee,
            roles=["Chủ tịch Hội đồng", "Thành viên", "Thành viên"],
            name_col=1, org_col=2, role_col=3,
        )
        session.save_close(doc)
    finally:
        session.quit()

    check = docx.Document(str(src))
    table = check.tables[0]
    assert table.cell(0, 0).text.strip() == "GS.TS. Nguyễn Công Khẩn"
    assert table.cell(0, 1).text.strip() == "Hội đồng Đạo đức Y sinh Quốc gia"
    assert table.cell(0, 2).text.strip() == "Chủ tịch Hội đồng"
    assert table.cell(1, 0).text.strip() == "PGS.TS.BS. Nguyễn Xuân Ninh"


def test_write_committee_secretaries_with_number_prefix(tmp_path):
    src = _table_fixture(tmp_path, rows=2, cols=2)
    session = word_writer.Session(force_backend="docx")
    try:
        doc = session.open(src)
        committee = _sample_committee()
        committee_writer.write_committee_secretaries(session, doc, 1, committee, number_prefix=True)
        session.save_close(doc)
    finally:
        session.quit()

    check = docx.Document(str(src))
    table = check.tables[0]
    assert table.cell(0, 0).text.strip() == "1. Hoàng Hà Linh"
    assert table.cell(0, 1).text.strip() == "Trung tâm NCKH - Viện VIAM"
    assert table.cell(1, 0).text.strip() == "2. Trương Phan Hồng Hà"


def test_write_committee_secretaries_without_prefix_has_no_degree(tmp_path):
    src = _table_fixture(tmp_path, rows=2, cols=2)
    session = word_writer.Session(force_backend="docx")
    try:
        doc = session.open(src)
        committee = _sample_committee()
        committee_writer.write_committee_secretaries(session, doc, 1, committee)
        session.save_close(doc)
    finally:
        session.quit()

    check = docx.Document(str(src))
    assert check.tables[0].cell(0, 0).text.strip() == "Hoàng Hà Linh"


def test_write_committee_roster_raises_on_role_count_mismatch(tmp_path):
    src = _table_fixture(tmp_path, rows=3, cols=3)
    session = word_writer.Session(force_backend="docx")
    try:
        doc = session.open(src)
        committee = _sample_committee()
        with pytest.raises(ValueError):
            committee_writer.write_committee_roster(
                session, doc, 1, committee, roles=["Chỉ một vai trò"]
            )
    finally:
        session.quit()


def test_roster_size_excludes_secretaries():
    committee = _sample_committee()
    assert committee_writer.roster_size(committee) == 3


def test_write_committee_roster_skips_org_column_when_org_col_is_none(tmp_path):
    src = _table_fixture(tmp_path, rows=3, cols=3)
    session = word_writer.Session(force_backend="docx")
    try:
        doc = session.open(src)
        committee = _sample_committee()
        committee_writer.write_committee_roster(
            session, doc, 1, committee,
            roles=["Chủ tịch Hội đồng", "Thành viên", "Thành viên"],
            name_col=1, org_col=None, role_col=None,
        )
        session.save_close(doc)
    finally:
        session.quit()

    check = docx.Document(str(src))
    table = check.tables[0]
    assert table.cell(0, 0).text.strip() == "GS.TS. Nguyễn Công Khẩn"
    assert table.cell(0, 1).text.strip() == ""


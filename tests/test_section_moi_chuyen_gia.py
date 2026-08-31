# test_section_moi_chuyen_gia.py
import shutil

import docx
import pytest

import excel_reader
import paths
import section_moi_chuyen_gia
import tokens
import word_writer
from excel_reader import CommitteeData, Person

CHECKLIST_PATH = paths.project_root() / excel_reader.CHECKLIST_FILENAME
SHEET_VIAM = "Đề tài - Bánh ăn dặm VIAM 2027"

MAU_DIR = paths.template_dir("03. Công văn mời chuyên gia - MẪU")
FILES = ["Công văn mời chuyên gia.docx", "Công văn mời chuyên gia nghiệm thu.docx"]


@pytest.fixture()
def dest_dir(tmp_path):
    for filename in FILES:
        shutil.copy2(MAU_DIR / filename, tmp_path / filename)
    return tmp_path


@pytest.fixture()
def info():
    return excel_reader.load_project_data(CHECKLIST_PATH, SHEET_VIAM)


def test_de_cuong_letter_has_one_page_per_external_member(dest_dir, info):
    session = word_writer.Session(force_backend="docx")
    try:
        section_moi_chuyen_gia.generate(session, dest_dir, info, tokens.build_common_tokens(info))
    finally:
        session.quit()

    expected = section_moi_chuyen_gia._dedupe_people(
        section_moi_chuyen_gia._external_members(info.ethics_committee, info.host_org)
        + section_moi_chuyen_gia._external_members(info.proposal_committee, info.host_org)
    )
    doc = docx.Document(str(dest_dir / "Công văn mời chuyên gia.docx"))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    for person in expected:
        assert f"{person.degree} {person.name}".strip() in full_text
    assert "{{CHUYEN_GIA_HO_TEN}}" not in full_text


def test_nghiem_thu_letter_only_pulls_from_acceptance_committee(dest_dir, info):
    session = word_writer.Session(force_backend="docx")
    try:
        section_moi_chuyen_gia.generate(session, dest_dir, info, tokens.build_common_tokens(info))
    finally:
        session.quit()

    doc = docx.Document(str(dest_dir / "Công văn mời chuyên gia nghiệm thu.docx"))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    expected = section_moi_chuyen_gia._external_members(info.acceptance_committee, info.host_org)
    for person in expected:
        assert f"{person.degree} {person.name}".strip() in full_text


def test_external_members_excludes_host_org_and_secretaries():
    committee = CommitteeData(
        chair=Person("Chủ tịch Ngoài", "PGS.", "Hội đồng ngoài"),
        reviewers=[Person("Phản biện VIAM", "TS.", "Viện Y học ứng dụng Việt Nam")],
        members=[Person("Ủy viên Ngoài", "ThS.", "Đơn vị khác")],
        secretaries=[Person("Thư ký VIAM", "CN.", "Ngoài")],
    )
    result = section_moi_chuyen_gia._external_members(committee, "Viện Y học ứng dụng Việt Nam")
    names = {p.name for p in result}
    assert names == {"Chủ tịch Ngoài", "Ủy viên Ngoài"}
    assert "Thư ký VIAM" not in names


def test_dedupe_people_collapses_same_person_across_committees():
    person_a = Person("Người A", "TS.", "Đơn vị A")
    person_a_dup = Person("người a", "TS.", "đơn vị a")
    person_b = Person("Người B", "TS.", "Đơn vị B")
    result = section_moi_chuyen_gia._dedupe_people([person_a, person_b, person_a_dup])
    assert result == [person_a, person_b]


def test_generate_prints_notice_when_no_external_members(tmp_path, info, capsys):
    import dataclasses

    for filename in FILES:
        shutil.copy2(MAU_DIR / filename, tmp_path / filename)

    all_host_committee = CommitteeData(
        chair=Person("Chủ tịch VIAM", "TS.", info.host_org),
        secretaries=[Person("Thư ký VIAM", "CN.", info.host_org)],
    )
    no_external_info = dataclasses.replace(
        info,
        ethics_committee=all_host_committee,
        proposal_committee=all_host_committee,
        acceptance_committee=all_host_committee,
    )

    session = word_writer.Session(force_backend="docx")
    try:
        section_moi_chuyen_gia.generate(
            session, tmp_path, no_external_info, tokens.build_common_tokens(no_external_info)
        )
    finally:
        session.quit()

    captured = capsys.readouterr()
    assert "bo qua dien thu moi de cuong" in captured.out
    assert "bo qua dien thu moi nghiem thu" in captured.out

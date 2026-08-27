import dataclasses
import shutil

import docx
import pytest

import excel_reader
import paths
import section_moi_chuyen_gia
import tokens
import word_writer

CHECKLIST_PATH = paths.project_root() / excel_reader.CHECKLIST_FILENAME
SHEET_VIAM = "Đề tài - Bánh ăn dặm VIAM 2027"

SOURCE_FILE = paths.project_root() / "03. Công văn mời chuyên gia - MẪU" / "Công văn mời chuyên gia.docx"


@pytest.fixture()
def dest_dir(tmp_path):
    shutil.copy2(SOURCE_FILE, tmp_path / SOURCE_FILE.name)
    return tmp_path


@pytest.fixture()
def info():
    return excel_reader.load_project_data(CHECKLIST_PATH, SHEET_VIAM)


def test_generate_replaces_title_and_removes_colostrum_intro(dest_dir, info):
    session = word_writer.Session(force_backend="docx")
    try:
        section_moi_chuyen_gia.generate(session, dest_dir, info, tokens.build_common_tokens(info))
    finally:
        session.quit()

    doc = docx.Document(str(dest_dir / "Công văn mời chuyên gia.docx"))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert info.title in full_text
    assert "LOF KUN COLOSTRUM" not in full_text
    assert f"01/{info.year} đến 12/{info.year}" in full_text


def test_generate_uses_parsed_timeline_not_just_year(dest_dir, info):
    custom_info = dataclasses.replace(info, timeline="Tháng 03/2027 đến tháng 09/2028")
    session = word_writer.Session(force_backend="docx")
    try:
        section_moi_chuyen_gia.generate(session, dest_dir, custom_info, tokens.build_common_tokens(custom_info))
    finally:
        session.quit()

    doc = docx.Document(str(dest_dir / "Công văn mời chuyên gia.docx"))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Thời gian thực hiện dự kiến: 03/2027 đến 09/2028." in full_text

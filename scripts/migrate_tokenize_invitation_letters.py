# migrate_tokenize_invitation_letters.py
import docx

import paths
from word_writer import _docx_replace_in_paragraph

MAU_DIR = paths.project_root() / "03. Công văn mời chuyên gia - MẪU"
DE_CUONG_PATH = MAU_DIR / "Công văn mời chuyên gia.docx"
NGHIEM_THU_PATH = MAU_DIR / "Công văn mời chuyên gia nghiệm thu.docx"

DE_CUONG_REPLACEMENTS = [
    (
        "Suy dinh dưỡng ở trẻ em dưới 5 tuổi – đặc biệt là suy dinh dưỡng thấp còi vẫn là một vấn đề sức khỏe cộng đồng. "
        "Một trong những giải pháp làm giảm tình trạng suy dinh dưỡng ở trẻ dưới 5 tuổi là sử dụng các sản phẩm bổ sung dinh dưỡng trong hệ thống trường mầm non. "
        "Nhằm đánh giá tình trạng suy dinh dưỡng ở trẻ dưới 5 tuổi và hiệu quả của sản phẩm bổ sung dinh dưỡng LOF KUN COLOSTRUM, Viện Y học ứng dụng Việt Nam tiến hành triển khai nghiên cứu",
        "[Bổ sung bối cảnh/lý do triển khai dự án tại đây] {{DON_VI_CHU_TRI}} tiến hành triển khai đề tài",
    ),
    (
        "Viện Y học ứng dụng Việt Nam trân trọng kính mời .................................................. tham gia Hội đồng khoa học xét duyệt đề cương và Hội đồng đạo đức nghiên cứu.",
        "{{DON_VI_CHU_TRI}} trân trọng kính mời {{CHUYEN_GIA_HO_TEN}} tham gia Hội đồng khoa học xét duyệt đề cương và Hội đồng đạo đức nghiên cứu.",
    ),
    (
        "Xin gửi hồ sơ nghiên cứu để .............................................................. đọc và cho ý kiến đóng góp, phản biện.",
        "Xin gửi hồ sơ nghiên cứu để {{CHUYEN_GIA_HO_TEN}} đọc và cho ý kiến đóng góp, phản biện.",
    ),
    (
        "Mọi thông tin chi tiết xin vui lòng liên hệ: Ông Lê Minh Khánh, Trung tâm nghiên cứu - Viện Y học ứng dụng Việt Nam (Email: leminhkhanh@viam.vn - Điện thoại: 096.3355.652). ",
        "Mọi thông tin chi tiết xin vui lòng liên hệ: {{DAU_MOI_LIEN_HE}}.",
    ),
]

NGHIEM_THU_REPLACEMENTS = [
    ("THAM GIA HỘI ĐỒNG KHOA HỌC VÀ HỘI ĐỒNG ĐẠO ĐỨC", "THAM GIA HỘI ĐỒNG NGHIỆM THU ĐỀ TÀI"),
    (
        "V/v: Mời chuyên gia tham gia Hội đồng khoa học và Hội đồng đạo đức",
        "V/v: Mời chuyên gia tham gia Hội đồng nghiệm thu đề tài",
    ),
    (
        "{{DON_VI_CHU_TRI}} trân trọng kính mời {{CHUYEN_GIA_HO_TEN}} tham gia Hội đồng khoa học xét duyệt đề cương và Hội đồng đạo đức nghiên cứu.",
        "{{DON_VI_CHU_TRI}} trân trọng kính mời {{CHUYEN_GIA_HO_TEN}} tham gia Hội đồng nghiệm thu đề tài.",
    ),
    (
        "Xin gửi hồ sơ nghiên cứu để {{CHUYEN_GIA_HO_TEN}} đọc và cho ý kiến đóng góp, phản biện.",
        "Xin gửi hồ sơ nghiệm thu để {{CHUYEN_GIA_HO_TEN}} đọc và cho ý kiến đánh giá, nghiệm thu.",
    ),
]

DOTS_REPLACEMENT = "{{CHUYEN_GIA_HO_TEN}}, {{CHUYEN_GIA_DON_VI}}."


def _is_dots_placeholder(text: str) -> bool:
    stripped = text.strip()
    return len(stripped) > 20 and set(stripped) <= {"."}


def _apply_replacements(doc, replacements) -> None:
    for paragraph in doc.paragraphs:
        for find, replace in replacements:
            _docx_replace_in_paragraph(paragraph, find, replace)
        if _is_dots_placeholder(paragraph.text):
            _docx_replace_in_paragraph(paragraph, paragraph.text, DOTS_REPLACEMENT)


def tokenize_de_cuong_letter() -> None:
    doc = docx.Document(str(DE_CUONG_PATH))
    _apply_replacements(doc, DE_CUONG_REPLACEMENTS)
    doc.save(str(DE_CUONG_PATH))


def create_nghiem_thu_letter() -> None:
    doc = docx.Document(str(DE_CUONG_PATH))
    _apply_replacements(doc, NGHIEM_THU_REPLACEMENTS)
    doc.save(str(NGHIEM_THU_PATH))


if __name__ == "__main__":
    tokenize_de_cuong_letter()
    create_nghiem_thu_letter()
    print("Da token hoa thu moi de cuong va tao file thu moi nghiem thu.")

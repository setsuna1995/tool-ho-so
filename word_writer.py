from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Optional

try:
    import win32com.client
    _WIN32COM_IMPORTED = True
except ImportError:
    _WIN32COM_IMPORTED = False

import docx

WD_FIND_WRAP_STOP = 0
WD_REPLACE_ALL = 2
WD_FORMAT_DOCX = 16


def _com_call(action, context: str):
    """Chạy `action` và bọc mọi exception với `context` để biết chính xác thao tác nào đang lỗi.

    Word COM chỉ trả về mã lỗi chung chung (vd "Command failed") không nói rõ
    đang thao tác ở file/ô nào - bọc lại giúp chẩn đoán nhanh hơn khi có lỗi.
    """
    try:
        return action()
    except Exception as e:
        raise RuntimeError(f"{context}: {e}") from e


def com_available() -> bool:
    if not _WIN32COM_IMPORTED:
        return False
    try:
        word = win32com.client.DispatchEx("Word.Application")
        word.Quit()
        return True
    except Exception:
        return False


@dataclass
class OpenDoc:
    backend: str
    handle: Any
    path: Path


def _docx_replace_in_paragraph(paragraph, find: str, replace: str) -> bool:
    full_text = "".join(run.text for run in paragraph.runs)
    if find not in full_text:
        return False
    new_text = full_text.replace(find, replace)
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""
    return True


def _docx_replace_text(document, find: str, replace: str, wildcards: bool) -> bool:
    search = find
    warn_paragraph_mark = False
    if wildcards and search.endswith("^13"):
        search = search[:-3]
        warn_paragraph_mark = True

    found = False
    for paragraph in document.paragraphs:
        if _docx_replace_in_paragraph(paragraph, search, replace):
            found = True

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if _docx_replace_in_paragraph(paragraph, search, replace):
                        found = True

    if found and warn_paragraph_mark:
        print(
            f"  [CANH BAO] Fallback docx: da xoa chu '{search.strip()}' nhung khong "
            "xoa duoc dong trong con lai (can Word COM moi lam duoc) - kiem tra thu cong."
        )

    return found


class Session:
    def __init__(self, force_backend: Optional[str] = None):
        self.backend = force_backend or ("com" if com_available() else "docx")
        self._word = None
        if self.backend == "com":
            self._word = win32com.client.DispatchEx("Word.Application")
            self._word.Visible = False

    def open(self, path: Path) -> OpenDoc:
        path = Path(path)
        if self.backend == "com":
            handle = _com_call(
                lambda: self._word.Documents.Open(str(path), False, False),
                f"Loi khi mo file '{path.name}' bang Word",
            )
        else:
            handle = docx.Document(str(path))
        return OpenDoc(backend=self.backend, handle=handle, path=path)

    def replace_text(
        self, doc: OpenDoc, find: str, replace: str, wildcards: bool = False, warn_if_missing: bool = True
    ) -> bool:
        if doc.backend == "com":
            def _do_find():
                rng = doc.handle.Content
                rng.Find.ClearFormatting()
                return bool(
                    rng.Find.Execute(
                        find, False, False, wildcards, False, False, True,
                        WD_FIND_WRAP_STOP, False, replace, WD_REPLACE_ALL,
                    )
                )

            preview_ctx = find if len(find) <= 60 else find[:60] + "..."
            found = _com_call(_do_find, f"Loi khi thay the '{preview_ctx}' trong file '{doc.path.name}'")
        else:
            found = _docx_replace_text(doc.handle, find, replace, wildcards)

        if not found and warn_if_missing:
            preview = find if len(find) <= 60 else find[:60] + "..."
            print(f"  [CANH BAO] Khong tim thay: '{preview}' trong {doc.path.name}")

        return found

    def replace_text_any(self, doc: OpenDoc, candidates: list, replace: str, wildcards: bool = False) -> bool:
        """Thử thay từng biến thể trong `candidates`, chỉ cảnh báo nếu KHÔNG biến thể nào khớp.

        Dùng cho các mẫu trộn lẫn nhiều cách viết hoa/thường của cùng một chỗ giữ
        chỗ (vd "20xx"/"20XX") - tránh cảnh báo giả khi backend Word COM đã thay
        cả hai kiểu chữ cùng lúc trong một lần tìm-thay không phân biệt hoa/thường.
        """
        found_any = False
        for find in candidates:
            if self.replace_text(doc, find, replace, wildcards=wildcards, warn_if_missing=False):
                found_any = True

        if not found_any:
            preview = ", ".join(f"'{c}'" for c in candidates)
            print(f"  [CANH BAO] Khong tim thay bat ky bien the nao trong [{preview}] trong {doc.path.name}")

        return found_any

    def fill_tokens(self, doc: OpenDoc, tokens: dict[str, str]) -> set[str]:
        """Ap dung moi token trong `tokens` vao `doc`, bo qua lang le token nao khong co mat.

        Khong phai template nao cung chua moi common token, nen luon goi voi
        warn_if_missing=False - mot token bi go sai/xoa nham se hien nguyen van
        '{{...}}' trong file .docx sinh ra, tu no da la dau hieu ro rang sai
        khi xem lai bang mat, khong can dua vao canh bao console.
        """
        return {token for token, value in tokens.items() if self.replace_text(doc, token, value, warn_if_missing=False)}

    def set_cell(self, doc: OpenDoc, table_index: int, row: int, col: int, text: str) -> None:
        if doc.backend == "com":
            _com_call(
                lambda: setattr(doc.handle.Tables.Item(table_index).Cell(row, col).Range, "Text", text),
                f"Loi khi ghi vao bang {table_index}, dong {row}, cot {col} trong file '{doc.path.name}'",
            )
        else:
            doc.handle.tables[table_index - 1].cell(row - 1, col - 1).text = text

    def save_close(self, doc: OpenDoc) -> None:
        if doc.backend == "com":
            def _do_save_close():
                doc.handle.Save()
                doc.handle.Close()

            _com_call(_do_save_close, f"Loi khi luu/dong file '{doc.path.name}'")
        else:
            doc.handle.save(str(doc.path))

    def quit(self) -> None:
        if self.backend == "com" and self._word is not None:
            self._word.Quit()
            self._word = None

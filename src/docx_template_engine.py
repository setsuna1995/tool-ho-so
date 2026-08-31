# docx_template_engine.py
from pathlib import Path
from typing import Any, Dict, Optional

import docx
import jinja2
from docxtpl import DocxTemplate


def _filter_upper(val: Any) -> str:
    if val is None:
        return ""
    return str(val).upper()


def _filter_lower(val: Any) -> str:
    if val is None:
        return ""
    return str(val).lower()


def _filter_title(val: Any) -> str:
    if val is None:
        return ""
    return str(val).title()


def _filter_default_placeholder(val: Any, default_val: str = "……………………") -> str:
    if val is None or not str(val).strip():
        return default_val
    return str(val)


def create_jinja_env() -> jinja2.Environment:
    """Tạo Jinja2 Environment với các custom filters hỗ trợ tiếng Việt & văn bản hành chính."""
    env = jinja2.Environment(autoescape=False)
    env.filters["upper"] = _filter_upper
    env.filters["lower"] = _filter_lower
    env.filters["title"] = _filter_title
    env.filters["default_placeholder"] = _filter_default_placeholder
    return env


def is_jinja_template(doc_path: Path) -> bool:
    """Kiểm tra xem file docx có chứa cú pháp Jinja2 (ví dụ: {{ project., {% for, {% if) hay không."""
    try:
        doc = docx.Document(str(doc_path))
        for p in doc.paragraphs:
            if "{{" in p.text or "{%" in p.text:
                return True
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if "{{" in p.text or "{%" in p.text:
                            return True
        return False
    except Exception:
        return False


def render_template(template_path: Path, output_path: Path, context: Dict[str, Any], env: Optional[jinja2.Environment] = None) -> None:
    """Render file docx mẫu bằng docxtpl với Jinja2 context và lưu vào output_path."""
    template_path = Path(template_path)
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    jinja_env = env or create_jinja_env()
    tpl = DocxTemplate(str(template_path))
    tpl.render(context, jinja_env=jinja_env)
    tpl.save(str(output_path))

import docx
import pytest

import word_writer


def _available_backends():
    backends = ["docx"]
    if word_writer.com_available():
        backends.append("com")
    return backends


BACKENDS = _available_backends()


def _make_paragraph_fixture(tmp_path, text):
    path = tmp_path / "sample.docx"
    d = docx.Document()
    d.add_paragraph(text)
    d.save(str(path))
    return path


def _make_table_fixture(tmp_path, rows=2, cols=2):
    path = tmp_path / "table.docx"
    d = docx.Document()
    d.add_table(rows=rows, cols=cols)
    d.save(str(path))
    return path


@pytest.mark.parametrize("backend", BACKENDS)
def test_replace_text_replaces_matching_text(tmp_path, backend):
    src = _make_paragraph_fixture(tmp_path, "Xin chao OLD_TITLE ban")

    session = word_writer.Session(force_backend=backend)
    try:
        doc = session.open(src)
        ok = session.replace_text(doc, "OLD_TITLE", "NEW_TITLE")
        assert ok is True
        session.save_close(doc)
    finally:
        session.quit()

    check = docx.Document(str(src))
    text = "\n".join(p.text for p in check.paragraphs)
    assert "NEW_TITLE" in text
    assert "OLD_TITLE" not in text


@pytest.mark.parametrize("backend", BACKENDS)
def test_replace_text_returns_false_when_not_found(tmp_path, backend):
    src = _make_paragraph_fixture(tmp_path, "Khong co gi lien quan")

    session = word_writer.Session(force_backend=backend)
    try:
        doc = session.open(src)
        ok = session.replace_text(doc, "KHONG_TON_TAI", "X")
        assert ok is False
        session.save_close(doc)
    finally:
        session.quit()


@pytest.mark.parametrize("backend", BACKENDS)
def test_set_cell_writes_value_into_table_cell(tmp_path, backend):
    src = _make_table_fixture(tmp_path)

    session = word_writer.Session(force_backend=backend)
    try:
        doc = session.open(src)
        session.set_cell(doc, 1, 1, 1, "GS.TS. Nguyen Cong Khan")
        session.save_close(doc)
    finally:
        session.quit()

    check = docx.Document(str(src))
    assert check.tables[0].cell(0, 0).text.strip() == "GS.TS. Nguyen Cong Khan"


def test_replace_text_wildcard_strips_matched_text_on_docx_backend(tmp_path):
    src = _make_paragraph_fixture(tmp_path, "Truoc. Cu nhan ABC. Sau.")

    session = word_writer.Session(force_backend="docx")
    try:
        doc = session.open(src)
        ok = session.replace_text(doc, "Cu nhan ABC", "", wildcards=True)
        assert ok is True
        session.save_close(doc)
    finally:
        session.quit()

    check = docx.Document(str(src))
    text = "\n".join(p.text for p in check.paragraphs)
    assert "Cu nhan ABC" not in text


def test_com_call_returns_action_result_on_success():
    assert word_writer._com_call(lambda: 42, "ctx") == 42


def test_com_call_wraps_exception_with_context():
    def boom():
        raise RuntimeError("boom goc")

    with pytest.raises(RuntimeError, match="ngu canh loi: boom goc"):
        word_writer._com_call(boom, "ngu canh loi")


def test_replace_text_any_does_not_warn_when_one_candidate_matches(tmp_path, capsys):
    src = _make_paragraph_fixture(tmp_path, "Nam 20xx la nam thuc hien")

    session = word_writer.Session(force_backend="docx")
    try:
        doc = session.open(src)
        ok = session.replace_text_any(doc, ["20XX", "20xx"], "2027")
        session.save_close(doc)
    finally:
        session.quit()

    captured = capsys.readouterr()
    assert "CANH BAO" not in captured.out
    assert ok is True

    check = docx.Document(str(src))
    text = "\n".join(p.text for p in check.paragraphs)
    assert "2027" in text


def test_replace_text_any_warns_once_when_no_candidate_matches(tmp_path, capsys):
    src = _make_paragraph_fixture(tmp_path, "Khong co gi lien quan")

    session = word_writer.Session(force_backend="docx")
    try:
        doc = session.open(src)
        ok = session.replace_text_any(doc, ["20XX", "20xx"], "2027")
        session.save_close(doc)
    finally:
        session.quit()

    captured = capsys.readouterr()
    assert captured.out.count("CANH BAO") == 1
    assert ok is False


def test_replace_text_wildcard_paragraph_mark_suffix_leaves_blank_paragraph(tmp_path, capsys):
    path = tmp_path / "wildcard_paragraph_mark.docx"
    d = docx.Document()
    d.add_paragraph("Truoc.")
    d.add_paragraph("Cu nhan ABC")
    d.add_paragraph("Sau.")
    d.save(str(path))

    session = word_writer.Session(force_backend="docx")
    try:
        doc = session.open(path)
        ok = session.replace_text(doc, "Cu nhan ABC^13", "", wildcards=True)
        assert ok is True
        session.save_close(doc)
    finally:
        session.quit()

    captured = capsys.readouterr()
    assert "CANH BAO" in captured.out

    check = docx.Document(str(path))
    assert len(check.paragraphs) == 3
    text = "\n".join(p.text for p in check.paragraphs)
    assert "Cu nhan ABC" not in text


def test_fill_tokens_replaces_every_present_token(tmp_path):
    src = _make_paragraph_fixture(tmp_path, "De tai {{TEN_DE_TAI}}, nam {{NAM}}.")

    session = word_writer.Session(force_backend="docx")
    try:
        doc = session.open(src)
        filled = session.fill_tokens(doc, {"{{TEN_DE_TAI}}": "ABC", "{{NAM}}": "2027"})
        session.save_close(doc)
    finally:
        session.quit()

    assert filled == {"{{TEN_DE_TAI}}", "{{NAM}}"}
    check = docx.Document(str(src))
    text = "\n".join(p.text for p in check.paragraphs)
    assert text == "De tai ABC, nam 2027."


def test_fill_tokens_silently_skips_absent_tokens_without_warning(tmp_path, capsys):
    src = _make_paragraph_fixture(tmp_path, "Chi co {{NAM}} o day.")

    session = word_writer.Session(force_backend="docx")
    try:
        doc = session.open(src)
        filled = session.fill_tokens(doc, {"{{NAM}}": "2027", "{{TEN_DE_TAI}}": "ABC"})
        session.save_close(doc)
    finally:
        session.quit()

    assert filled == {"{{NAM}}"}
    captured = capsys.readouterr()
    assert "CANH BAO" not in captured.out

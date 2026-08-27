import shutil
from pathlib import Path

import docx
import pytest

import migrate_templates_to_tokens as migrate
import word_writer

# apply_mapping BAT BUOC dung backend COM (xem chu thich trong ham do), nen ca
# hai test duoi day can Word tren may. Theo dung quy uoc san co cua repo cho
# test phu thuoc COM - xem test_convert_doc_templates.py.
pytestmark = pytest.mark.skipif(
    not word_writer.com_available(), reason="Can Word COM de chay test nay"
)


def test_apply_mapping_replaces_old_text_with_token(tmp_path):
    src = tmp_path / "sample.docx"
    d = docx.Document()
    d.add_paragraph("Tên đề tài: OLD_TITLE_MARKER.")
    d.save(str(src))

    migrate.apply_mapping(src, [("OLD_TITLE_MARKER", "{{TEN_DE_TAI}}")])

    check = docx.Document(str(src))
    text = "\n".join(p.text for p in check.paragraphs)
    assert text == "Tên đề tài: {{TEN_DE_TAI}}."


def test_apply_mapping_raises_when_search_text_not_found(tmp_path):
    src = tmp_path / "sample.docx"
    d = docx.Document()
    d.add_paragraph("Không liên quan.")
    d.save(str(src))

    try:
        migrate.apply_mapping(src, [("KHONG_TON_TAI", "{{X}}")])
        assert False, "Ky vong RuntimeError khi khong tim thay chuoi can thay"
    except RuntimeError as e:
        assert "KHONG_TON_TAI" in str(e)

# test_migrate_tokenize_invitation_letters.py
import docx

import migrate_tokenize_invitation_letters as migrate


def test_de_cuong_letter_has_no_leftover_sample_text():
    migrate.tokenize_de_cuong_letter()

    doc = docx.Document(str(migrate.DE_CUONG_PATH))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "LOF KUN COLOSTRUM" not in full_text
    assert "Lê Minh Khánh" not in full_text
    assert "{{CHUYEN_GIA_HO_TEN}}" in full_text
    assert "{{CHUYEN_GIA_DON_VI}}" in full_text
    assert "{{DAU_MOI_LIEN_HE}}" in full_text


def test_nghiem_thu_letter_created_with_nghiem_thu_wording():
    migrate.tokenize_de_cuong_letter()
    migrate.create_nghiem_thu_letter()

    doc = docx.Document(str(migrate.NGHIEM_THU_PATH))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "NGHIỆM THU" in full_text
    assert "HỘI ĐỒNG KHOA HỌC VÀ HỘI ĐỒNG ĐẠO ĐỨC" not in full_text
    assert "{{CHUYEN_GIA_HO_TEN}}" in full_text
    assert "{{DAU_MOI_LIEN_HE}}" in full_text


def test_de_cuong_tokenization_is_idempotent():
    migrate.tokenize_de_cuong_letter()
    migrate.tokenize_de_cuong_letter()

    doc = docx.Document(str(migrate.DE_CUONG_PATH))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert full_text.count("{{CHUYEN_GIA_HO_TEN}}") == 3

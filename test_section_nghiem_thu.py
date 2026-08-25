import dataclasses
import shutil

import docx
import pytest

import excel_reader
import paths
import section_nghiem_thu
import word_writer

CHECKLIST_PATH = paths.project_root() / excel_reader.CHECKLIST_FILENAME
SHEET_VIAM = "Đề tài - Bánh ăn dặm VIAM 2027"

SOURCE_DIR = paths.project_root() / "04. Hồ sơ nghiệm thu - MẪU"
FILES = [
    "9. Quyết định thành lập HĐ nghiệm thu.docx",
    "10. Biên bản họp HĐ nghiệm thu.docx",
    "11. Biên bản kiểm phiếu nghiệm thu.docx",
    "12. Quyết định công nhận kết quả đề tài.docx",
    "Phiếu chấm điểm nghiệm thu (TVCT_ĐGHQ).docx",
    "Phiếu chấm điểm nghiệm thu (TNLS).docx",
    "Phiếu ký nhận tiền.docx",
    "Phiếu nhận xét nghiệm thu.docx",
]


@pytest.fixture()
def dest_dir(tmp_path):
    for filename in FILES:
        shutil.copy2(SOURCE_DIR / filename, tmp_path / filename)
    return tmp_path


@pytest.fixture()
def info():
    return excel_reader.load_project_data(CHECKLIST_PATH, SHEET_VIAM)


def test_generate_writes_correct_secretary_org(dest_dir, info):
    session = word_writer.Session(force_backend="docx")
    try:
        section_nghiem_thu.generate(session, dest_dir, info)
    finally:
        session.quit()

    doc = docx.Document(str(dest_dir / "9. Quyết định thành lập HĐ nghiệm thu.docx"))
    secretary_table = doc.tables[2]
    assert secretary_table.cell(0, 0).text.strip() == "1. Hoàng Hà Linh"
    assert secretary_table.cell(0, 1).text.strip() == "Viện Y học Ứng dụng Việt Nam"


def test_generate_uses_dynamic_member_count_not_hardcoded_05(dest_dir, info):
    session = word_writer.Session(force_backend="docx")
    try:
        section_nghiem_thu.generate(session, dest_dir, info)
    finally:
        session.quit()

    doc = docx.Document(str(dest_dir / "10. Biên bản họp HĐ nghiệm thu.docx"))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "là 05 người" in full_text


def test_generate_replaces_both_year_cases_without_false_warning(dest_dir, info, capsys):
    session = word_writer.Session(force_backend="docx")
    try:
        section_nghiem_thu.generate(session, dest_dir, info)
    finally:
        session.quit()

    captured = capsys.readouterr()
    assert "Khong tim thay bat ky bien the nao" not in captured.out

    doc = docx.Document(str(dest_dir / "12. Quyết định công nhận kết quả đề tài.docx"))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text += "\n" + cell.text
    assert "20xx" not in full_text
    assert "20XX" not in full_text
    assert str(info.year) in full_text


def test_generate_keeps_tvct_and_removes_tnls_by_default(dest_dir, info):
    session = word_writer.Session(force_backend="docx")
    try:
        section_nghiem_thu.generate(session, dest_dir, info)
    finally:
        session.quit()

    assert (dest_dir / "Phiếu chấm điểm nghiệm thu (TVCT_ĐGHQ).docx").exists()
    assert not (dest_dir / "Phiếu chấm điểm nghiệm thu (TNLS).docx").exists()


def test_generate_selects_tnls_form_and_removes_tvct_when_research_type_is_tnls(dest_dir, info):
    tnls_info = dataclasses.replace(info, research_type="TNLS")
    session = word_writer.Session(force_backend="docx")
    try:
        section_nghiem_thu.generate(session, dest_dir, tnls_info)
    finally:
        session.quit()

    doc = docx.Document(str(dest_dir / "Phiếu chấm điểm nghiệm thu (TNLS).docx"))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert tnls_info.title in full_text
    assert not (dest_dir / "Phiếu chấm điểm nghiệm thu (TVCT_ĐGHQ).docx").exists()


def test_generate_raises_for_unrecognized_research_type(dest_dir, info):
    bad_info = dataclasses.replace(info, research_type="LOAI_LA")
    session = word_writer.Session(force_backend="docx")
    try:
        with pytest.raises(ValueError):
            section_nghiem_thu.generate(session, dest_dir, bad_info)
    finally:
        session.quit()

import dataclasses
import shutil

import docx
import pytest

import excel_reader
import paths
import section_khoa_hoc
import tokens
import word_writer

CHECKLIST_PATH = paths.project_root() / excel_reader.CHECKLIST_FILENAME
SHEET_VIAM = "Đề tài - Bánh ăn dặm VIAM 2027"

SOURCE_DIR = paths.project_root() / "02. Hồ sơ khoa học đề cương - MẪU"
FILES = [
    "05. QĐ TLHĐ khoa học xét đề cương.docx",
    "06. BB họp thông qua đề cương.docx",
    "07. BB kiểm phiếu thông qua đề cương.docx",
    "08. QĐ phê duyệt đề tài.docx",
    "Phiếu chấm điểm HĐ đề cương.docx",
    "Phiếu nhận xét đánh giá hồ sơ.docx",
]


@pytest.fixture()
def dest_dir(tmp_path):
    for filename in FILES:
        shutil.copy2(SOURCE_DIR / filename, tmp_path / filename)
    return tmp_path


@pytest.fixture()
def info():
    return excel_reader.load_project_data(CHECKLIST_PATH, SHEET_VIAM)


def test_generate_fixes_proposal_secretary_org(dest_dir, info):
    session = word_writer.Session(force_backend="docx")
    try:
        section_khoa_hoc.generate(session, dest_dir, info, tokens.build_common_tokens(info))
    finally:
        session.quit()

    doc = docx.Document(str(dest_dir / "05. QĐ TLHĐ khoa học xét đề cương.docx"))
    secretary_table = doc.tables[2]
    assert secretary_table.cell(0, 1).text.strip() == "Hoàng Hà Linh"
    assert secretary_table.cell(0, 2).text.strip() == "Trung tâm NCKH - Viện VIAM"


def test_generate_uses_parsed_timeline_not_just_year(dest_dir, info):
    custom_timeline = "Tháng 03/2027 đến tháng 09/2028"
    start, end = excel_reader.parse_timeline(custom_timeline)
    custom_info = dataclasses.replace(
        info,
        timeline=custom_timeline,
        common_tokens={
            **info.common_tokens,
            "{{THOI_GIAN_BAT_DAU}}": start,
            "{{THOI_GIAN_KET_THUC}}": end,
        },
    )
    session = word_writer.Session(force_backend="docx")
    try:
        section_khoa_hoc.generate(session, dest_dir, custom_info, tokens.build_common_tokens(custom_info))
    finally:
        session.quit()

    doc = docx.Document(str(dest_dir / "08. QĐ phê duyệt đề tài.docx"))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "từ tháng 03/2027 đến tháng 09/2028" in full_text
